# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches
import re
import os


regex_letter = re.compile(r"^(Oficio\s+)([0-9])(.*)$")
regex_date = re.compile(r"(Rio Pardo de Minas\/MG, )([0-9]{2})( de )(.*)(de 2024)")
regex_contract = re.compile(r"(contrato nº )([0-9]{1,})")
regex_code = re.compile(r"(código )([0-9]{1,})")
regex_item = re.compile(r"(item )([0-9]{1,})")
regex_company = re.compile(r"^Empresa:\s+(.*?)(?=\s*CNPJ:)")

def get_data_from_doc(doc_name):
    try:
        with open(doc_name, "rb") as doc_file:
            doc = Document(doc_file)
    except FileNotFoundError:
        print(f"O arquivo '{doc_name}' não foi encontrado.")
        return None
    
    letter_number = ''
    letter_date = ''
    contract = ''
    code = ''
    item = ''
    company = ''

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if match := regex_letter.search(text):
            letter_number = match.group(0) 
        if match := regex_date.search(text):
            letter_date = match.group(0)[23:] 
        if match := regex_contract.search(text):
            contract = match.group(0)  
        if match := regex_code.search(text):
            code = match.group(0) 
        if match := regex_item.search(text):
            item = match.group(0) 
        if match := regex_company.search(text):
            company = match.group(1) 

    if (len(code) == 0):
        code = item
    return {
        'letter_number': letter_number,
        'letter_date': letter_date,
        'contract_number': contract,
        'item_code': code,
        'company_name': company
    }
    

def get_letters_name(folder):
    files = []
    for file in os.listdir(folder):
        if file.endswith(".docx") or file.endswith(".doc"):
            files.append(file)
    return files


def create_word_table(data_array, template_path):
    # Abre o documento de template
    doc = Document(template_path)

    # Procura o texto "[TABELA]" no documento
    for paragraph in doc.paragraphs:
        if paragraph.text == '[TABELA]':
            # Remove o parágrafo com o texto "[TABELA]"
            paragraph.clear()

            # Cria uma tabela com 5 colunas e 1 linha (para o cabeçalho)
            # table = doc.add_table(rows=1, cols=5, style='Table Grid')
            table = doc.add_table(rows=1, cols=5)

            # Define os cabeçalhos da tabela
            table.cell(0, 0).text = 'Ofício'
            table.cell(0, 1).text = 'Data'
            table.cell(0, 2).text = 'Contrato'
            table.cell(0, 3).text = 'Item'
            table.cell(0, 4).text = 'Empresa'

            # Itera sobre o array de dados e preenche a tabela
            for data in data_array:
                row = table.add_row().cells
                row[0].text = data.get('oficio_name', '')
                row[1].text = data.get('oficio_date', '')
                row[2].text = data.get('contract_number', '')
                row[3].text = data.get('item_code', '')
                row[4].text = data.get('company_name', '')

            # Define o estilo da tabela
            # style = doc.styles['Table Grid']
            # font = style.font
            # font.name = 'Times New Roman'
            # font.size = Pt(12)

            break
    doc.save("pareceres/Parecer_unico_aditivos.docx")

def main():
    folder = 'of_aditivos'
    letters = get_letters_name(folder)
    letters_data = []
    for doc_name in letters:
        letters_data.append(get_data_from_doc(f'{folder}/{doc_name}'))

    template_path = 'modelos/Parecer_aditivo.docx'
    create_word_table(letters_data, template_path)

if __name__ == "__main__":
   main()