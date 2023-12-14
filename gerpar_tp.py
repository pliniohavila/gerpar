# -*- coding: utf-8 -*-
from docx import Document
from openpyxl import load_workbook
import datetime

def data_para_string(data):
    if (type(data) == str):
        return data
    return data.strftime("%d/%m/%Y")
    
def parecer_edital(assunto, processo, num_modalidade, data):
    with open("modelos/tomada_de_preco_edital.docx", "rb") as docFile:
        doc = Document(docFile)

    for paragraph in doc.paragraphs:
        if "[ASSUNTO]" in paragraph.text:
            for run in paragraph.runs:
                if "[ASSUNTO]" in run.text:
                    run.text = run.text.replace("[ASSUNTO]", assunto)
                    break

        elif "[MODALIDADE_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[MODALIDADE_N]", num_modalidade)
            paragraph.text = new_text

        elif "[PROCESSO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PROCESSO_N]", processo)
            paragraph.text = new_text
      
        elif "[DATA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[DATA]", data_para_string(data))
            paragraph.text = new_text

    nome_documento = str.replace(processo, '/', '-')
    doc.save("pareceres/Parecer_tp_Edital_%s.docx" % (nome_documento))


def parecer_contrato(assunto, processo, num_modalidade, data):
    with open("modelos/tomada_de_preco_contrato.docx", "rb") as docFile:
        doc = Document(docFile)

    for paragraph in doc.paragraphs:
        if "[ASSUNTO]" in paragraph.text:
            for run in paragraph.runs:
                if "[ASSUNTO]" in run.text:
                    run.text = run.text.replace("[ASSUNTO]", assunto)
                    break
        elif "[MODALIDADE_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[MODALIDADE_N]", num_modalidade)
            paragraph.text = new_text

        elif "[PROCESSO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PROCESSO_N]", processo)
            paragraph.text = new_text

        elif "[DATA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[DATA]", data_para_string(data))
            paragraph.text = new_text

    nome_documento = str.replace(processo, '/', '-')
    doc.save("pareceres/Parecer_tp_Contrato_%s.docx" % (nome_documento))

def make_parecer(data):
    assunto = data[0]
    processo = data[1]
    num_modalidade = data[2]
    dataEdital = data[3]
    dataContrato = data[4]

    parecer_edital(assunto, processo, num_modalidade, dataEdital)
    parecer_contrato(assunto, processo, num_modalidade, dataContrato)


def main():
    wb = load_workbook(filename='relacao_contratos_tp.xlsx')
    ws = wb.active
    for row in ws.iter_rows(values_only=True):
        make_parecer(row)

if __name__ == "__main__":
   main()
