# -*- coding: utf-8 -*-
from docx import Document

# Para criar o parecer do edital
def data_para_string(data):
    if (type(data) == str):
        return data
    return data.strftime("%d/%m/%Y")
    
def parecer_edital(assunto, processo, num_modalidade, data):

    # Abre o documento modelo de base
    with open("modelos/pregao_edital.docx", "rb") as docFile:
        doc = Document(docFile)

    for paragraph in doc.paragraphs:

        if "[ASSUNTO]" in paragraph.text:
            for run in paragraph.runs:
                if "[ASSUNTO]" in run.text:
                    run.text = run.text.replace("[ASSUNTO]", assunto)
                    break

        elif "[MODALIDADE_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[MODALIDADE_N]", num_modalidade)
            paragraph.text = new_text

        elif "[PROCESSO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PROCESSO_N]", processo)
            paragraph.text = new_text
      
        elif "[DATA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[DATA]", data_para_string(data))
            paragraph.text = new_text

    # Salva o novo parecer criado
    nome_documento = str.replace(processo, '/', '-')
    doc.save("pareceres/Parecer_Edital_%s.docx" % (nome_documento))

def parecer_contrato(assunto, processo, num_modalidade, data):
    # Abre o documento modelo de base
    with open("modelos/pregao_contrato.docx", "rb") as docFile:
        doc = Document(docFile)

    for paragraph in doc.paragraphs:
        if "[ASSUNTO]" in paragraph.text:
            for run in paragraph.runs:
                if "[ASSUNTO]" in run.text:
                    run.text = run.text.replace("[ASSUNTO]", assunto)
                    break
        elif "[MODALIDADE_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[MODALIDADE_N]", num_modalidade)
            paragraph.text = new_text

        elif "[PROCESSO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PROCESSO_N]", processo)
            paragraph.text = new_text

        elif "[DATA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[DATA]", data_para_string(data))
            paragraph.text = new_text

    # Salva o novo parecer criado
    nome_documento = str.replace(processo, '/', '-')
    doc.save("pareceres/Parecer_Contrato_%s.docx" % (nome_documento))


def parecer_contrato_prorrogacao_servicos_conntinuados(tipo, contrato, processo, contratada, cnpj, objeto, prazo_prorrogacao, data):

    if tipo == 1:
        with open("modelos/prorrogacao_contrato_serviços_continuados.docx", "rb") as docFile:
            doc = Document(docFile)
    if tipo == 2:
        with open("modelos/prorrogacao_contrato_locacao.docx", "rb") as docFile:
            doc = Document(docFile)
    if tipo == 4:
        with open("modelos/prorrogacao_termo_credenciamento_serviços_continuados.docx", "rb") as docFile:
            doc = Document(docFile)

    for paragraph in doc.paragraphs:
        if "[CONTRATO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[CONTRATO_N]", contrato)
            paragraph.text = new_text

        elif "[PROCESSO_N]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PROCESSO_N]", processo)
            paragraph.text = new_text

        elif "[CONTRATADA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[CONTRATADA]", contratada)
            paragraph.text = new_text

        elif "[CNPJ]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[CNPJ]", cnpj)
            paragraph.text = new_text

        elif "[OBJETO]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[OBJETO]", objeto)
            paragraph.text = new_text

        elif "[PRAZO_PRORROGACAO]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[PRAZO_PRORROGACAO]", data_para_string(prazo_prorrogacao))
            paragraph.text = new_text

        elif "[DATA]" in paragraph.text:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, "[DATA]", data_para_string(data))
            paragraph.text = new_text

    # Salva o novo parecer criado
    nome_documento = str.replace(contrato, '/', '-')
    if tipo == 4:
        contratada = str.replace(contratada, '/', '-')
        contratada = str.replace(contratada, ' ', '_')
        doc.save("pareceres/Parecer_Prorrogacao_-_Credenciamento_%s_%s.docx" % (nome_documento, contratada))
    else:
        doc.save("pareceres/Parecer_Prorrogacao_-_Contrato_%s.docx" % (nome_documento))    
